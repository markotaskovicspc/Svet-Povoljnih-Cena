from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "outputs/Uputstvo-za-Google-Facebook-Apple-login.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 33, 46)
MUTED = RGBColor(90, 98, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_GOLD = "FFF7DF"
BORDER = "B8C2CC"
BLACK = RGBColor(0, 0, 0)

CALLBACK_URL = "https://vyebjbcfhgujlvjnoxpl.supabase.co/auth/v1/callback"
SUPABASE_DOMAIN = "vyebjbcfhgujlvjnoxpl.supabase.co"
PROJECT_REF = "vyebjbcfhgujlvjnoxpl"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_element(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    element.get_or_add_pPr().append(shd)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, total_dxa=None, indent_dxa=120):
    total_dxa = total_dxa or sum(widths_dxa)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    for old_grid in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    for tag, attrs, text in [
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {}, " PAGE "),
        ("w:fldChar", {"w:fldCharType": "separate"}, None),
        ("w:t", {}, "1"),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ]:
        run = paragraph.add_run()
        element = OxmlElement(tag)
        for key, value in attrs.items():
            element.set(qn(key), value)
        if text is not None:
            element.text = text
        run._r.append(element)


def create_numbering(doc, kind="decimal"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_numbered_list(doc, items):
    num_id = create_numbering(doc, "decimal")
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, num_id)
        add_rich_text(p, item)


def add_bullet_list(doc, items):
    num_id = create_numbering(doc, "bullet")
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, num_id)
        add_rich_text(p, item)


def add_rich_text(paragraph, content):
    if isinstance(content, str):
        run = paragraph.add_run(content)
        set_run_font(run)
        return
    for part in content:
        text = part.get("text", "")
        run = paragraph.add_run(text)
        set_run_font(
            run,
            name=part.get("font", "Calibri"),
            size=part.get("size"),
            color=part.get("color"),
            bold=part.get("bold"),
            italic=part.get("italic"),
        )


def paragraph(text="", style=None, before=None, after=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        add_rich_text(p, text)
    return p


def code_paragraph(text):
    p = doc.add_paragraph(style="CodeValue")
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(5)
    shade_element(p._p, LIGHT_GRAY)
    run = p.add_run(text)
    set_run_font(run, name="Courier New", size=9.5, color=INK)
    return p


def add_note(title, body, fill=PALE_GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="D9C16C", size="6")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    r = cell.paragraphs[0].add_run(title)
    set_run_font(r, size=10.5, bold=True, color=INK)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(body), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_label_table(rows, widths=(2700, 6660), header=None):
    offset = 1 if header else 0
    table = doc.add_table(rows=len(rows) + offset, cols=2)
    set_table_geometry(table, list(widths), indent_dxa=120)
    set_table_borders(table)
    if header:
        cell = table.cell(0, 0)
        cell.merge(table.cell(0, 1))
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, bold=True, color=INK)
    for idx, (label, value) in enumerate(rows, start=offset):
        label_cell = table.cell(idx, 0)
        value_cell = table.cell(idx, 1)
        shade_cell(label_cell, LIGHT_GRAY)
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        set_run_font(lp.add_run(label), bold=True, color=INK)
        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        if isinstance(value, list):
            add_rich_text(vp, value)
        else:
            set_run_font(vp.add_run(value), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_source_item(label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_hyperlink(p, label, url)


def configure_styles(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.25
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.25
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.color.rgb = DARK_BLUE
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.25
    h3.paragraph_format.keep_with_next = True

    for name, base in (("DocTitle", "Normal"), ("DocSubtitle", "Normal"), ("CodeValue", "Normal")):
        if name not in styles:
            styles.add_style(name, 1)
    title = styles["DocTitle"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(7)
    title.paragraph_format.line_spacing = 1.15

    subtitle = styles["DocSubtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.25

    code = styles["CodeValue"]
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(9.5)
    code.font.color.rgb = INK
    code.paragraph_format.left_indent = Inches(0.08)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.line_spacing = 1.15

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    set_run_font(footer_p.add_run("Social login uputstvo | Strana "), size=9, color=MUTED)
    add_page_field(footer_p)


doc = Document()
configure_styles(doc)

p = paragraph("Uputstvo za podešavanje Google, Facebook i Apple login-a", style="DocTitle")
p.paragraph_format.keep_with_next = True
paragraph(
    "Jednostavni koraci za klijenta: gde da klikne, koje URL adrese da unese i koje podatke da pošalje developeru.",
    style="DocSubtitle",
)
add_label_table(
    [
        ("Projekat", "Svet Akcija - Supabase Auth"),
        ("Supabase project ref", PROJECT_REF),
        ("Verzija dokumenta", "23.06.2026."),
        ("Stil rada", "Klijent pravi naloge/aplikacije, developer lepi dobijene podatke u Supabase."),
    ],
    widths=(2600, 6760),
)

doc.add_heading("Pre početka", level=1)
paragraph(
    "Ovaj dokument je za osobu koja ima pristup nalozima firme. Nije potrebno menjati kod sajta. Cilj je da se naprave Google, Facebook i Apple aplikacije za prijavu korisnika i da se developeru pošalju potrebni ID/secret podaci."
)
add_note(
    "Važno za bezbednost",
    "Client Secret, App Secret, Apple .p8 fajl i slične podatke ne slati javno, u screenshotovima, niti u grupama gde ima više ljudi. Najbolje je poslati ih developeru kroz dogovoreni siguran kanal.",
)

doc.add_page_break()
doc.add_heading("Podaci koji se koriste u sva tri podešavanja", level=2)
add_label_table(
    [
        ("Supabase callback URL", CALLBACK_URL),
        ("Supabase domen za Apple", SUPABASE_DOMAIN),
        (
            "Domen sajta",
            "Upisati pravi domen sajta. Primeri: https://svetpovoljnihcena.rs i, ako postoji, https://www.svetpovoljnihcena.rs",
        ),
        ("Lokalno testiranje", "Samo za developera: http://localhost:3000"),
    ],
    widths=(2600, 6760),
    header="Tačne vrednosti za kopiranje",
)

doc.add_heading("1. Google login", level=1)
paragraph("Ovo se podešava u Google Cloud Console.")
add_numbered_list(
    doc,
    [
        "Otvorite https://console.cloud.google.com/ i ulogujte se na Google nalog firme.",
        "Gore levo izaberite postojeći projekat za sajt. Ako projekat ne postoji, kliknite New Project i napravite novi projekat.",
        "U levom meniju otvorite Google Auth Platform.",
        "Ako Google traži osnovno podešavanje aplikacije, popunite App name, User support email i Developer contact email.",
        "U delu Audience izaberite External, osim ako login smeju da koriste samo ljudi iz iste Google Workspace organizacije.",
        "U delu Data Access / Scopes proverite da postoje osnovni scope-ovi: openid, email i profile.",
        "U levom meniju kliknite Clients.",
        "Kliknite Create client.",
        "Za Application type izaberite Web application.",
        "Za Name unesite npr. Svet Akcija Web Login.",
        "U Authorized JavaScript origins kliknite Add URI i unesite domen sajta, bez dodatnih putanja.",
        "U Authorized redirect URIs kliknite Add URI i unesite Supabase callback URL iz ovog dokumenta.",
        "Kliknite Create.",
        "Google će prikazati Client ID i Client Secret. Sačuvajte oba podatka i pošaljite ih developeru.",
    ],
)
doc.add_heading("Google - vrednosti za unos", level=2)
add_label_table(
    [
        ("Authorized JavaScript origins", "https://svetpovoljnihcena.rs"),
        ("Dodatni origin ako postoji www", "https://www.svetpovoljnihcena.rs"),
        ("Lokalni origin za developera", "http://localhost:3000"),
        ("Authorized redirect URI", CALLBACK_URL),
    ],
    widths=(3100, 6260),
)
add_note(
    "Napomena",
    "Ako je pravi domen drugačiji od primera svetpovoljnihcena.rs, unesite stvarni domen. U Authorized JavaScript origins ide samo početak adrese, npr. https://example.com, bez /login ili drugih putanja.",
    fill=LIGHT_BLUE,
)

doc.add_heading("2. Facebook login", level=1)
paragraph("Ovo se podešava u Meta for Developers, ranije Facebook Developers.")
add_numbered_list(
    doc,
    [
        "Otvorite https://developers.facebook.com/ i ulogujte se na Facebook nalog koji ima pristup firmi.",
        "Gore desno kliknite My Apps.",
        "Kliknite Create App.",
        "Ako Meta pita za use case, izaberite opciju za Authentication / Facebook Login. Ako traži tip aplikacije, izaberite potrošačku/web aplikaciju koja omogućava Facebook login.",
        "Unesite ime aplikacije, npr. Svet Akcija Login, i kontakt email firme.",
        "Kada se otvori dashboard aplikacije, pronađite Facebook Login i kliknite Set up. Ako postoji Quickstart, preskočite ga.",
        "U levom meniju otvorite Facebook Login > Settings.",
        "U polje Valid OAuth Redirect URIs nalepite Supabase callback URL iz ovog dokumenta.",
        "Kliknite Save Changes.",
        "U levom meniju otvorite Use Cases.",
        "Pronađite Authentication and Account Creation i kliknite Edit.",
        "Proverite da public_profile i email postoje i da su spremni za testiranje. Ako email nije dodat, kliknite Add pored email.",
        "U levom meniju otvorite Settings > Basic.",
        "U App Domains unesite domen sajta bez https://, npr. svetpovoljnihcena.rs. Ako koristite www, dodajte i www.svetpovoljnihcena.rs.",
        "Dodajte Privacy Policy URL, kategoriju i ostala obavezna polja koja Meta traži.",
        "U Settings > Basic kopirajte App ID i App Secret. App Secret se prikazuje klikom na Show.",
        "Pošaljite developeru App ID i App Secret.",
        "Kada je sve testirano, aplikaciju treba prebaciti u Live mode da bi login radili svi korisnici, a ne samo test korisnici.",
    ],
)
doc.add_heading("Facebook - vrednosti za unos", level=2)
add_label_table(
    [
        ("Valid OAuth Redirect URIs", CALLBACK_URL),
        ("App Domains", "svetpovoljnihcena.rs, www.svetpovoljnihcena.rs"),
        ("Permissions", "public_profile i email"),
        ("Podaci za developera", "App ID i App Secret"),
    ],
    widths=(3100, 6260),
)

doc.add_heading("3. Apple login", level=1)
paragraph(
    "Apple login je najzahtevniji deo. Potreban je Apple Developer nalog. Ako firma nema Apple Developer Program, Apple login se ne može potpuno podesiti dok se taj nalog ne napravi i aktivira."
)
add_note(
    "Važno za Apple",
    "Apple client secret ističe i mora da se obnavlja. Za OAuth podešavanje ga treba rotirati najkasnije na svakih 6 meseci. .p8 fajl obavezno čuvati na sigurnom mestu.",
)

doc.add_heading("3.1 Napravite Primary App ID", level=2)
add_numbered_list(
    doc,
    [
        "Otvorite https://developer.apple.com/account/ i ulogujte se na Apple Developer nalog firme.",
        "Otvorite Certificates, Identifiers & Profiles.",
        "U levom meniju kliknite Identifiers.",
        "Kliknite dugme plus (+).",
        "Izaberite App IDs, zatim App, pa Continue.",
        "U Description unesite npr. Svet Akcija Primary App ID.",
        "Za Bundle ID izaberite Explicit i unesite dogovoreni identifikator, npr. rs.svetpovoljnihcena.app.",
        "U listi Capabilities uključite Sign in with Apple.",
        "Kliknite Continue, zatim Register.",
    ],
)

doc.add_heading("3.2 Napravite Services ID za web login", level=2)
add_numbered_list(
    doc,
    [
        "U Identifiers kliknite plus (+).",
        "Izaberite Services IDs i kliknite Continue.",
        "U Description unesite npr. Svet Akcija Web Login.",
        "U Identifier unesite dogovoreni identifikator, npr. rs.svetpovoljnihcena.web.",
        "Kliknite Continue, zatim Register.",
        "Vratite se na listu Identifiers i u filteru izaberite Services IDs.",
        "Otvorite Services ID koji ste upravo napravili.",
        "Uključite Sign in with Apple i kliknite Configure.",
        "U Primary App ID izaberite App ID koji ste napravili u prethodnom delu.",
        "U Domains and Subdomains unesite Supabase domen iz ovog dokumenta.",
        "U Return URLs unesite Supabase callback URL iz ovog dokumenta.",
        "Kliknite Done, zatim Continue ili Save, zavisno od toga šta Apple prikaže.",
    ],
)
doc.add_heading("Apple - vrednosti za Services ID", level=2)
add_label_table(
    [
        ("Domains and Subdomains", SUPABASE_DOMAIN),
        ("Return URLs", CALLBACK_URL),
        ("Services ID primer", "rs.svetpovoljnihcena.web"),
        ("Primary App ID primer", "rs.svetpovoljnihcena.app"),
    ],
    widths=(3100, 6260),
)

doc.add_page_break()
doc.add_heading("3.3 Napravite Apple private key", level=2)
add_numbered_list(
    doc,
    [
        "U Certificates, Identifiers & Profiles otvorite Keys.",
        "Kliknite plus (+).",
        "Unesite ime ključa, npr. Svet Akcija Sign in with Apple Key.",
        "Uključite Sign in with Apple.",
        "Kliknite Configure i izaberite Primary App ID koji ste napravili.",
        "Kliknite Save, zatim Continue, zatim Register.",
        "Kliknite Download i preuzmite .p8 fajl. Ovaj fajl se može preuzeti samo jednom.",
        "Zapišite Key ID koji Apple prikaže.",
        "Zapišite Team ID. Najčešće se vidi u gornjem desnom meniju Apple Developer naloga ili u Membership delu.",
    ],
)
doc.add_heading("Apple - šta poslati developeru", level=2)
add_bullet_list(
    doc,
    [
        "Team ID",
        "Services ID, npr. rs.svetpovoljnihcena.web",
        "Key ID",
        ".p8 fajl koji je Apple preuzeo",
    ],
)
paragraph(
    "Developer od ovih podataka generiše Apple Client Secret i unosi ga u Supabase. Klijent ne mora ručno da pravi JWT ako nije tehnički obučen."
)

doc.add_heading("4. Unos podataka u Supabase", level=1)
paragraph("Ovaj deo najčešće radi developer, ali može ga uraditi i osoba koja ima pristup Supabase projektu.")
add_numbered_list(
    doc,
    [
        "Otvorite https://supabase.com/dashboard i izaberite projekat.",
        "U levom meniju kliknite Authentication.",
        "Otvorite Sign In / Providers.",
        "Otvorite Google, uključite provider i unesite Google Client ID i Client Secret.",
        "Otvorite Facebook, uključite provider i unesite Facebook App ID kao Client ID i Facebook App Secret kao Client Secret.",
        "Otvorite Apple, uključite provider, unesite Apple Services ID kao Client ID i unesite Apple Client Secret koji je developer generisao.",
        "Kliknite Save posle svakog providera.",
    ],
)
doc.add_heading("Supabase URL Configuration", level=2)
add_numbered_list(
    doc,
    [
        "U Supabase-u otvorite Authentication > URL Configuration.",
        "U Site URL unesite produkcioni domen sajta, npr. https://svetpovoljnihcena.rs.",
        "U Redirect URLs dodajte produkcione URL adrese koje aplikacija koristi posle login-a. Za developer testiranje može se dodati i http://localhost:3000/**.",
        "Sačuvajte izmene.",
    ],
)

doc.add_heading("5. Završna provera", level=1)
add_bullet_list(
    doc,
    [
        "Testirati Google login u privatnom/incognito prozoru.",
        "Testirati Facebook login u privatnom/incognito prozoru.",
        "Testirati Apple login ako je Apple Developer podešavanje završeno.",
        "Proveriti da se korisnik posle login-a vraća na sajt, a ne na Supabase ili stranicu sa greškom.",
        "Ako login prijavi redirect_uri_mismatch ili sličnu grešku, skoro uvek je problem pogrešno upisan callback URL.",
    ],
)
add_note(
    "Kratko pravilo",
    "Za Google i Facebook se u provider dashboard najvažnije unosi isti Supabase callback URL. Za Apple se kao domen koristi Supabase domen, a kao Return URL isti Supabase callback URL.",
    fill=LIGHT_BLUE,
)

doc.add_heading("Izvori", level=1)
paragraph("Dokument je pripremljen prema zvaničnoj dokumentaciji i Supabase uputstvima za social login.")
add_source_item("Supabase - Google login", "https://supabase.com/docs/guides/auth/social-login/auth-google")
add_source_item("Supabase - Facebook login", "https://supabase.com/docs/guides/auth/social-login/auth-facebook")
add_source_item("Supabase - Apple login", "https://supabase.com/docs/guides/auth/social-login/auth-apple")
add_source_item("Supabase - Redirect URLs", "https://supabase.com/docs/guides/auth/redirect-urls")
add_source_item("Google OAuth for web applications", "https://developers.google.com/identity/protocols/oauth2/web-server")
add_source_item("Meta - Facebook Login for the Web", "https://developers.facebook.com/docs/facebook-login/web")
add_source_item("Apple - Configure Sign in with Apple for the web", "https://developer.apple.com/help/account/capabilities/configure-sign-in-with-apple-for-the-web")
add_source_item("Apple - Create a Sign in with Apple private key", "https://developer.apple.com/help/account/capabilities/create-a-sign-in-with-apple-private-key")

doc.core_properties.title = "Uputstvo za podešavanje Google, Facebook i Apple login-a"
doc.core_properties.subject = "Social login setup za Supabase projekat"
doc.core_properties.author = ""
doc.core_properties.comments = "Generated for client handoff."
doc.save(OUT)
print(OUT)
